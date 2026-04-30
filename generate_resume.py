from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Page margins (tight for 1 page) ──
for section in doc.sections:
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)

# ── Helpers ──
def set_font(run, bold=False, size=10, color=None):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(1)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '7B68EE')
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_section_title(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(text.upper())
    set_font(run, bold=True, size=10, color=(75, 0, 130))
    add_divider(doc)

def add_bullet(doc, text, bold_prefix=None, size=9.5):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.left_indent = Inches(0.15)
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        set_font(r1, bold=True, size=size)
        r2 = p.add_run(text)
        set_font(r2, size=size)
    else:
        r = p.add_run(text)
        set_font(r, size=size)

def add_job_header(doc, company, dates, title, location):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(0)
    r1 = p.add_run(company)
    set_font(r1, bold=True, size=10)
    r2 = p.add_run(f"  {dates}")
    set_font(r2, size=9.5, color=(100,100,100))
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(1)
    r3 = p2.add_run(title)
    set_font(r3, bold=True, size=9.5, color=(75,0,130))
    r4 = p2.add_run(f"  |  {location}")
    set_font(r4, size=9.5, color=(100,100,100))

# ══════════════════════════════════════════
# NAME & CONTACT
# ══════════════════════════════════════════
p_name = doc.add_paragraph()
p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_name.paragraph_format.space_before = Pt(0)
p_name.paragraph_format.space_after = Pt(2)
r = p_name.add_run("MANSI KAMOTHI")
set_font(r, bold=True, size=16, color=(75,0,130))

p_contact = doc.add_paragraph()
p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_contact.paragraph_format.space_before = Pt(0)
p_contact.paragraph_format.space_after = Pt(3)
r2 = p_contact.add_run(
    "San Francisco, CA  |  (510) 710-8030  |  mansikamothi1999@gmail.com  |  LinkedIn  |  GitHub  |  Portfolio"
)
set_font(r2, size=9.5, color=(80,80,80))

# ══════════════════════════════════════════
# PROFESSIONAL SUMMARY
# ══════════════════════════════════════════
add_section_title(doc, "Professional Summary")
p_sum = doc.add_paragraph()
p_sum.paragraph_format.space_before = Pt(1)
p_sum.paragraph_format.space_after = Pt(1)
r_sum = p_sum.add_run(
    "Product Analyst with 3+ years in SaaS delivering $1.2M ARR growth, 19.8% conversion lift, and $420K in cost savings. "
    "Experienced in AI/ML tooling (Salesforce Agentforce, Einstein Analytics, OpenAI API), Zendesk support analytics, "
    "UX-informed product instrumentation, SQL, Python, Tableau, and Salesforce CRM. "
    "Certified: Google PM · AWS Cloud Practitioner · Salesforce AI Associate & Agentforce Specialist."
)
set_font(r_sum, size=9.5)

# ══════════════════════════════════════════
# CORE COMPETENCIES (single-line format)
# ══════════════════════════════════════════
add_section_title(doc, "Core Competencies")

competency_lines = [
    ("Analytics & Data: ", "SQL (Advanced), Python (pandas, numpy, scikit-learn), BigQuery, Snowflake, Statistical Analysis, Forecasting, Churn Prediction"),
    ("AI & Automation: ", "Salesforce Agentforce, Einstein Analytics, OpenAI API, LLM reporting automation, Predictive lead scoring, Anomaly detection"),
    ("BI & Visualization: ", "Tableau, Power BI, Looker Studio, Amplitude, GA4, Mixpanel, Optimizely"),
    ("UX & Product: ", "User journey mapping, Funnel & usability analysis, Heatmap/click-path analysis, Feature instrumentation, A/B experimentation"),
    ("Platforms & Tools: ", "Salesforce CRM, Zendesk, Stripe, Airflow, Segment, GCP, JIRA, Confluence, Agile/Scrum"),
    ("Certifications: ", "Google Project Management · AWS Cloud Practitioner · Salesforce AI Associate & Agentforce Specialist"),
]
for bold_part, rest in competency_lines:
    add_bullet(doc, rest, bold_prefix=bold_part, size=9.5)

# ══════════════════════════════════════════
# PROFESSIONAL EXPERIENCE
# ══════════════════════════════════════════
add_section_title(doc, "Professional Experience")

# ── Tungsten Automation ──
add_job_header(doc, "Tungsten Automation", "Jul 2024 – Present", "Product Analyst (Product Management Focus)", "Remote, USA")

tungsten_bullets = [
    "Increased conversion 19.8% and generated $1.2M ARR growth by analysing onboarding funnels and implementing A/B experimentation frameworks.",
    "Analysed 45M+ monthly product events using SQL and Snowflake across Salesforce, Stripe, and Zendesk, uncovering 3 growth opportunities and reducing churn by 4%.",
    "Leveraged Salesforce Agentforce to automate support routing and case classification, reducing resolution time ~30% and freeing CS capacity for high-value interactions.",
    "Applied AI-assisted anomaly detection (Python/scikit-learn) on product events to proactively flag churn signals, enabling CS intervention before cancellation.",
    "Integrated OpenAI API into reporting workflows to auto-generate executive narrative summaries from dashboards, reducing report prep time by 40%.",
    "Analyzed Zendesk ticket trends and CSAT data using SQL, enabling the product team to prioritize 3 bug fixes that reduced support ticket volume by 18%.",
    "Built Zendesk + Salesforce + Stripe integrated pipelines providing a single source of truth for customer health scoring across leadership.",
    "Partnered with UX and Product Design to instrument user flows in Amplitude and GA4; identified 4 onboarding drop-off points that, when redesigned, contributed to the 19.8% conversion lift.",
    "Designed KPI frameworks and executive dashboards used by 120+ stakeholders; automated pipelines eliminating 85 hrs/week of manual reporting.",
    "Identified operational inefficiencies through data analysis, reducing costs by $420K annually.",
]
for b in tungsten_bullets:
    add_bullet(doc, b)

# ── Magna Infotech ──
add_job_header(doc, "Magna Infotech", "Mar 2020 – Dec 2021", "Business Analyst, Product", "India")

magna_bullets = [
    "Increased checkout conversion 14% and feature adoption 22% through A/B testing, funnel optimisation, and customer behaviour analysis.",
    "Collaborated with UX and Product Design to map user journeys from event data, surfacing friction points that informed UI improvements.",
    "Built Tableau and Power BI dashboards visualising product KPIs, revenue performance, and engagement trends for cross-functional teams.",
    "Automated analytics workflows reducing reporting time 25%; defined KPI frameworks tracked via Google Analytics and Salesforce CRM.",
    "Converted business requirements into technical specifications, reducing development rework and accelerating feature delivery.",
]
for b in magna_bullets:
    add_bullet(doc, b)

# ══════════════════════════════════════════
# EDUCATION
# ══════════════════════════════════════════
add_section_title(doc, "Education")
p_edu = doc.add_paragraph()
p_edu.paragraph_format.space_before = Pt(2)
p_edu.paragraph_format.space_after = Pt(0)
r_e1 = p_edu.add_run("California State University, Sacramento")
set_font(r_e1, bold=True, size=10)
r_e2 = p_edu.add_run("  |  Master of Science, Electrical & Electronics Engineering  |  Jan 2022 – Dec 2024")
set_font(r_e2, size=9.5, color=(80,80,80))

# ══════════════════════════════════════════
# SAVE
# ══════════════════════════════════════════
output_path = "/Users/mansi/Desktop/portfolio/Mansi_Kamothi_Resume_Updated.docx"
doc.save(output_path)
print(f"Saved: {output_path}")